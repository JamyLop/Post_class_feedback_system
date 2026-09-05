"""一生一案导出 · 正式数据版式

不沿用旧 4 列窄表，改为贴合当前正式模型的分节结构：
  封面 / 档案信息(含入学成绩、家长、健康) / 总览(总体问题·升学目标·当前状态)
  / 学科方案(每科：诊断归因·目标强化·任务清单·执行记录·督查) / 附录(版本/说明)

版式目标：A4、2cm 边距、微软雅黑、可扫读、表头浅色、长文本不挤压。
"""

from __future__ import annotations

from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.student_case import CaseCycle, CaseReview, CaseTask, StudentCase, SubjectPlan, TaskCheckin

SUBJECT_ORDER = ["语文", "数学", "英语", "物理", "化学", "生物", "政治", "历史", "地理", "德育"]

STATUS_LABEL = {
    "draft": "草稿",
    "pending_confirmation": "待德育审查",
    "revision_required": "需整改",
    "executing": "执行中",
    "pending_review": "待复盘",
    "adjusted": "已调整",
    "archived": "已归档",
}

# 品牌色：与前端一致的克制配色
NAVY = RGBColor(0x1B, 0x2F, 0x4A)
NAVY_LIGHT = "1B2F4A"
NAVY_HEX = "1B2F4A"
ACCENT = RGBColor(0x2F, 0x6F, 0xED)
MUTED = RGBColor(0x64, 0x74, 0x8B)
BORDER_COLOR = "E2E8F0"
HEADER_FILL = "F1F5F9"
HEADER_FILL_DARK = "1E3A5F"
LIGHT_FILL = "F8FAFC"

TEMPLATE_COVER = Path(__file__).resolve().parents[3] / "docx" / "一生一案" / "导出模板" / "封面.docx"
# 用户确认的正式导出底稿。缺失时才退回旧版程序化版式，避免部署漏带模板时导出中断。
CASE_EXPORT_TEMPLATE = Path(__file__).resolve().parents[3] / "docx" / "一生一案" / "导出模板" / "一生一案_导出模板.docx"
COVER_IMAGE = Path(__file__).resolve().parents[3] / "docx" / "一生一案" / "导出模板" / "学生阶段发展规划封面.png"


def _prepend_cover_image(doc: Document) -> None:
    """把学校封面作为独立的 A4 零边距分节，正文继续沿用模板边距。"""
    if not COVER_IMAGE.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    # 强制按 A4 成品尺寸拉伸，确保图片覆盖整页，不留下上下或左右白边。
    paragraph.add_run().add_picture(str(COVER_IMAGE), width=Cm(21.0), height=Cm(29.7))
    body = doc._body._element
    body.remove(paragraph._p)
    body.insert(0, paragraph._p)

    # 段落级 sectPr 结束封面分节，并从下一页开始恢复模板正文分节。
    section_break = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = OxmlElement("w:sectPr")
    section_type = OxmlElement("w:type")
    section_type.set(qn("w:val"), "nextPage")
    sect_pr.append(section_type)
    page_size = OxmlElement("w:pgSz")
    page_size.set(qn("w:w"), "11906")
    page_size.set(qn("w:h"), "16838")
    sect_pr.append(page_size)
    page_margins = OxmlElement("w:pgMar")
    for margin in ("top", "right", "bottom", "left"):
        page_margins.set(qn(f"w:{margin}"), "0")
    sect_pr.append(page_margins)
    p_pr.append(sect_pr)
    section_break.append(p_pr)
    body.insert(1, section_break)


def _format_date(d) -> str:
    if not d:
        return ""
    try:
        return d.strftime("%Y/%m/%d")
    except Exception:
        return str(d).strip()


def _cadence_label(c: str) -> str:
    return {"daily": "日", "weekly": "周", "monthly": "月"}.get(c, c or "")


def _status_label(s: str) -> str:
    return STATUS_LABEL.get(s, s or "")


def _set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for k, v in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{k}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_table_borders(tbl, color=BORDER_COLOR, size="6") -> None:
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._element.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _set_fonts(run, size=9, bold=False, color=None, name="微软雅黑", east_asia="微软雅黑") -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:hint"), "eastAsia")


def _add_para(doc, text, size=9, bold=False, color=None, align=None, space_before=0, space_after=4, line_spacing=1.15, name="微软雅黑"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.widowControl = True
    if align is not None:
        p.alignment = align
    run = p.add_run(text or "")
    _set_fonts(run, size=size, bold=bold, color=color, name=name)
    return p


def _add_mixed_para(doc, parts: list[tuple[str, dict]], align=None, space_before=0, space_after=4):
    """parts: [(text, {size,bold,color,name})]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    for text, style in parts:
        run = p.add_run(text)
        _set_fonts(run, size=style.get("size", 9), bold=style.get("bold", False), color=style.get("color"), name=style.get("name", "微软雅黑"))
    return p


def _split_overview(text: str) -> list[tuple[str, str]]:
    """按；或;拆分，提取 学科：内容"""
    if not text or not text.strip():
        return []
    import re
    # 统一分隔
    raw = re.split(r"[；;]\s*", text.strip())
    out: list[tuple[str, str]] = []
    for seg in raw:
        if not seg.strip():
            continue
        seg = seg.strip()
        # 尝试按第一个：或:分割
        m = re.match(r"^([^：:]{1,12})[：:]\s*(.+)$", seg, re.S)
        if m:
            out.append((m.group(1).strip(), m.group(2).strip()))
        else:
            # 无学科前缀，作为整体
            out.append(("", seg))
    return out


def _add_section_heading(doc, num: str, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keepWithNext = True
    # 左侧竖线用编号+标题模拟
    r1 = p.add_run(f"{num}  ")
    _set_fonts(r1, size=8, bold=True, color=ACCENT)
    r2 = p.add_run(title)
    _set_fonts(r2, size=13, bold=True, color=NAVY)
    if subtitle:
        r3 = p.add_run(f"  ·  {subtitle}")
        _set_fonts(r3, size=8, bold=False, color=MUTED)
    # 下划线
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E2E8F0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_sub_heading(doc, title: str, desc: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("▎ ")
    _set_fonts(r1, size=9, bold=True, color=ACCENT)
    r2 = p.add_run(title)
    _set_fonts(r2, size=10, bold=True, color=NAVY)
    if desc:
        r3 = p.add_run(f"  {desc}")
        _set_fonts(r3, size=7.5, bold=False, color=MUTED)


def _ensure_a4_margins(doc: Document) -> None:
    for sec in doc.sections:
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.9)
        sec.right_margin = Cm(1.9)
        sec.header_distance = Cm(1)
        sec.footer_distance = Cm(1)


def _add_header_footer(doc: Document, cycle_name: str = "") -> None:
    # 页眉：学校 + 周期
    sec = doc.sections[0]
    header = sec.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("白粉高级英才中学  ·  一生一案  " + (f"·  {cycle_name}" if cycle_name else ""))
    _set_fonts(run, size=7, bold=False, color=MUTED)
    # 页脚：页码
    footer = sec.footer
    footer.is_linked_to_previous = False
    p2 = footer.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p2._p.get_or_add_pPr()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("w:space"), "preserve"); instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    # 用简单文本代替复杂域，避免兼容性，改用静态提示由 Word 自动更新
    run2 = p2.add_run()
    rPr = run2._element.get_or_add_rPr()
    fldChar1_ = OxmlElement("w:fldChar"); fldChar1_.set(qn("w:fldCharType"), "begin")
    run2._element.append(fldChar1_)
    run2 = p2.add_run()
    run2.text = " PAGE "
    _set_fonts(run2, size=7, color=MUTED)
    # 简化：不依赖域，直接写 — 第 X 页 —
    p2.text = ""
    run3 = p2.add_run("—  第  页  —")
    _set_fonts(run3, size=7, color=MUTED)


def _add_cover(doc: Document, student_name: str, class_name: str, cycle_name: str, case: StudentCase, profile: Any | None = None) -> None:
    # 顶部装饰线
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(36)
    p0.paragraph_format.space_after = Pt(6)
    r0 = p0.add_run("白粉高级英才中学")
    _set_fonts(r0, size=10, bold=False, color=MUTED, name="微软雅黑")
    # 字母间距通过字符间距实现（w:spacing）
    r0._element.get_or_add_rPr().append(_spacing_el(80))

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run("一生一案")
    _set_fonts(r1, size=28, bold=True, color=NAVY, name="微软雅黑")
    r1._element.get_or_add_rPr().append(_spacing_el(120))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("个性化学习发展指导方案")
    _set_fonts(r2, size=11, bold=False, color=MUTED)
    r2._element.get_or_add_rPr().append(_spacing_el(60))

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(18)
    # 装饰线：用段落边框
    pPr = p3._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    r3 = p3.add_run("  —  ONE STUDENT ONE PLAN  —  ")
    _set_fonts(r3, size=7, bold=False, color=RGBColor(0x94, 0xA3, 0xB8))
    # 尝试嵌入封面图（若存在），作为装饰而非依赖版式
    try:
        if TEMPLATE_COVER.exists():
            import zipfile
            with zipfile.ZipFile(str(TEMPLATE_COVER)) as z:
                if "word/media/image1.png" in z.namelist():
                    data = z.read("word/media/image1.png")
                    tmp = Path(BytesIO(data).getvalue())  # not used
                    # 直接用 BytesIO 写入
                    bio = BytesIO(data)
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_after = Pt(12)
                    run_img = p_img.add_run()
                    run_img.add_picture(bio, width=Cm(3.2))
    except Exception:
        pass

    # 学生姓名大字
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(8)
    p_name.paragraph_format.space_after = Pt(6)
    rn = p_name.add_run(student_name or "—")
    _set_fonts(rn, size=26, bold=True, color=NAVY)
    rn._element.get_or_add_rPr().append(_spacing_el(40))

    # 班级 / 学年 / 版本 / 状态 一行信息卡
    info = [
        class_name or "—",
        cycle_name or "—",
        f"V{case.version}  {_status_label(case.status)}",
        _format_date(getattr(case, "updated_at", None)) or _format_date(datetime.now(timezone.utc)),
    ]
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(22)
    for i, txt in enumerate(info):
        r = p_info.add_run(txt)
        _set_fonts(r, size=8, bold=False, color=MUTED)
        if i < len(info) - 1:
            sep = p_info.add_run("   ·   ")
            _set_fonts(sep, size=8, color=RGBColor(0xCB, 0xD5, 0xE1))

    # 信息卡片表格（4 列）
    tbl = doc.add_table(rows=2, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    _set_table_borders(tbl, color=BORDER_COLOR, size="4")
    headers = ["班级", "学年周期", "版本 / 状态", "建档日期"]
    vals = [
        class_name or "—",
        cycle_name or "—",
        f"V{case.version}  ·  {_status_label(case.status)}",
        _format_date(case.created_at) if getattr(case, "created_at", None) else _format_date(datetime.now(timezone.utc)),
    ]
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, HEADER_FILL)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_fonts(run, size=7.5, bold=True, color=MUTED)
    for ci, v in enumerate(vals):
        cell = tbl.cell(1, ci)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell, top=80, bottom=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(v)
        _set_fonts(run, size=9, bold=True, color=NAVY)
    # 列宽
    for row in tbl.rows:
        for idx, w in enumerate([Cm(4.2), Cm(4.8), Cm(4.5), Cm(3.5)]):
            row.cells[idx].width = w

    # 底部说明
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_before = Pt(18)
    p_note.paragraph_format.space_after = Pt(2)
    rn2 = p_note.add_run("本方案由班主任牵头、学科教师协同制定，德育审查后进入执行；执行中持续记录任务与督查，阶段复盘生成新版本。")
    _set_fonts(rn2, size=7, color=RGBColor(0x94, 0xA3, 0xB8))
    p_note2 = doc.add_paragraph()
    p_note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn3 = p_note2.add_run(f"导出时间：{datetime.now(timezone.utc).strftime('%Y/%m/%d %H:%M')}  ·  系统自动生成，仅供校内使用")
    _set_fonts(rn3, size=6.5, color=RGBColor(0x94, 0xA3, 0xB8))


def _spacing_el(val: int) -> OxmlElement:
    el = OxmlElement("w:spacing")
    el.set(qn("w:val"), str(val))
    return el


def _pv(profile: Any | None, key: str):
    if profile is None:
        return None
    if isinstance(profile, dict):
        return profile.get(key)
    return getattr(profile, key, None)


def _add_profile_section(doc: Document, profile: Any | None, guardians: list | None = None) -> None:
    _add_section_heading(doc, "01", "学生档案", "Student Profile")
    if not profile:
        _add_para(doc, "尚未建立学生档案。", size=8.5, color=MUTED, space_before=2)
        return

    def val(key: str) -> str:
        v = _pv(profile, key)
        if v is None:
            return "—"
        s = str(v).strip()
        return s if s else "—"

    def fmt_scores() -> None:
        # 入学成绩卡片
        has_any = any(_pv(profile, f) is not None for f in [
            "entrance_total_score","entrance_chinese","entrance_math","entrance_english",
            "entrance_physics","entrance_chemistry","entrance_biology","entrance_politics","entrance_history","entrance_geography"
        ])
        if not has_any and not val("entrance_scores").strip().replace("—",""):
            _add_para(doc, "入学成绩：暂未录入", size=8.5, color=MUTED, space_before=4, space_after=2)
            return
        tbl = doc.add_table(rows=2, cols=6)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(tbl, color=BORDER_COLOR)
        headers = ["总分", "语文", "数学", "英语", "物理", "化学"]
        headers2 = ["生物", "政治", "历史", "地理", "—", "—"]
        # 第一行
        keys1 = ["entrance_total_score","entrance_chinese","entrance_math","entrance_english","entrance_physics","entrance_chemistry"]
        keys2 = ["entrance_biology","entrance_politics","entrance_history","entrance_geography", None, None]
        for ci, h in enumerate(headers):
            cell = tbl.cell(0, ci)
            _set_cell_shading(cell, HEADER_FILL); _set_cell_margins(cell, top=50, bottom=50)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h); _set_fonts(run, size=7, bold=True, color=MUTED)
        for ci, k in enumerate(keys1):
            cell = tbl.cell(1, ci)
            _set_cell_margins(cell, top=70, bottom=70)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            v = _pv(profile, k) if k else None
            txt = str(v) if v is not None else "—"
            run = p.add_run(txt); _set_fonts(run, size=9, bold=True, color=NAVY)
        # 第二行（追加一行视觉上合并，改用第二张表更简洁：这里用 1 张 3 行模拟）
        tbl2 = doc.add_table(rows=2, cols=6)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(tbl2, color=BORDER_COLOR)
        for ci, h in enumerate(["生物","政治","历史","地理","备注",""]):
            cell = tbl2.cell(0, ci)
            _set_cell_shading(cell, HEADER_FILL); _set_cell_margins(cell, top=50, bottom=50)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if h:
                run = p.add_run(h); _set_fonts(run, size=7, bold=True, color=MUTED)
        for ci, k in enumerate(keys2):
            cell = tbl2.cell(1, ci)
            _set_cell_margins(cell, top=60, bottom=60)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if k is None:
                if ci == 4:
                    txt = val("entrance_scores") if val("entrance_scores") != "—" else "—"
                    run = p.add_run(txt[:30]); _set_fonts(run, size=7, color=MUTED)
                continue
            v = _pv(profile, k)
            txt = str(v) if v is not None else "—"
            run = p.add_run(txt); _set_fonts(run, size=9, bold=True, color=NAVY)

    # 1) 学生信息 3×4 网格
    tbl = doc.add_table(rows=3, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(tbl, color=BORDER_COLOR)
    grid = [
        [("姓名", val("student_name")), ("性别", val("gender")), ("民族", val("ethnicity")), ("年级", val("grade"))],
        [("生源学校", val("source_school")), None, ("家长", f"{val('parent_name')}（{val('parent_relationship')}）"), ("电话", val("parent_phone"))],
        [("家长评价", val("parent_evaluation")), None, None, None],
    ]
    # 合并处理：第二行生源学校跨2列，第三行家长评价跨全行
    # 先填充
    for r, row in enumerate(grid):
        if r == 0:
            for c, pair in enumerate(row):
                label, value = pair
                cell = tbl.cell(r, c)
                _set_cell_shading(cell, HEADER_FILL) if c % 2 == 0 else None
                # 标签列浅色
                if c % 2 == 0:
                    _set_cell_shading(cell, HEADER_FILL)
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                if c % 2 == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(label); _set_fonts(run, size=7.5, bold=True, color=MUTED)
                else:
                    run = p.add_run(value); _set_fonts(run, size=8.5, bold=False, color=NAVY)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif r == 1:
            # 生源学校跨2列：合并 cell(1,0) 与 (1,1) 的值
            c0 = tbl.cell(r, 0); _set_cell_shading(c0, HEADER_FILL); _set_cell_margins(c0)
            p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run("生源学校"); _set_fonts(run, size=7.5, bold=True, color=MUTED)
            # 合并视觉：把 0,1 两格内容放在 1 列，另一列留空但合并边框
            c1 = tbl.cell(r, 1); _set_cell_margins(c1)
            p = c1.paragraphs[0]; run = p.add_run(val("source_school")); _set_fonts(run, size=8.5, color=NAVY)
            # 家长
            c2 = tbl.cell(r, 2); _set_cell_shading(c2, HEADER_FILL); _set_cell_margins(c2)
            p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run("家长"); _set_fonts(run, size=7.5, bold=True, color=MUTED)
            c3 = tbl.cell(r, 3); _set_cell_margins(c3)
            p = c3.paragraphs[0]; run = p.add_run(f"{val('parent_name')}  {val('parent_phone')}  {val('parent_relationship')}"); _set_fonts(run, size=8, color=NAVY)
        else:
            # 家长评价 / 主要需求 合并
            c0 = tbl.cell(r, 0); _set_cell_shading(c0, HEADER_FILL); _set_cell_margins(c0)
            p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run("家庭反馈"); _set_fonts(run, size=7.5, bold=True, color=MUTED)
            # 跨3列合并：用 _merge
            # python-docx 合并
            try:
                tbl.cell(r, 1).merge(tbl.cell(r, 3))
            except Exception:
                pass
            c1 = tbl.cell(r, 1); _set_cell_margins(c1)
            p = c1.paragraphs[0]
            txt = f"家长评价：{val('parent_evaluation')}    |    主要需求：{val('primary_needs')}"
            run = p.add_run(txt); _set_fonts(run, size=8, color=NAVY)
    # 列宽
    widths = [Cm(2.6), Cm(4.2), Cm(2.6), Cm(5.0)]
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    _add_sub_heading(doc, "入学成绩", "Entrance Scores")
    fmt_scores()

    # 健康信息（仅当 health_visible 或需要提示）
    hv = _pv(profile, "health_visible")
    if hv is None:
        hv = True
    if hv is False:
        _add_para(doc, "健康与体检信息已设为仅校长可见，此处不展示具体内容。", size=7.5, color=MUTED, space_before=4)
    else:
        has_health = any([val("allergy_history") != "—", val("underlying_conditions") != "—", val("other_health_notes") != "—"])
        if has_health:
            _add_sub_heading(doc, "健康与体检", "Health Notes")
            health_tbl = doc.add_table(rows=1, cols=3)
            health_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(health_tbl)
            headers = ["过敏史", "隐性疾病", "其他"]
            keys = ["allergy_history", "underlying_conditions", "other_health_notes"]
            # header row
            for ci, h in enumerate(headers):
                cell = health_tbl.cell(0, ci)
                _set_cell_shading(cell, HEADER_FILL); _set_cell_margins(cell)
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h); _set_fonts(run, size=7, bold=True, color=MUTED)
            row = health_tbl.add_row()
            for ci, k in enumerate(keys):
                cell = row.cells[ci]
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                run = p.add_run(val(k)); _set_fonts(run, size=8, color=NAVY)


def _add_overview_section(doc: Document, case: StudentCase) -> None:
    _add_section_heading(doc, "02", "总览", "Overview")
    overall = (case.overall_problem or "").strip()
    target = (case.admission_target or "").strip()
    summary = (case.current_summary or "").strip()

    # 总体问题
    _add_sub_heading(doc, "总体问题", "按学科归纳的诊断摘要")
    if not overall:
        _add_para(doc, "尚未填写总体问题。", size=8.5, color=MUTED)
    else:
        items = _split_overview(overall)
        if len(items) == 1 and not items[0][0]:
            _add_para(doc, items[0][1], size=9, color=NAVY, space_after=2)
        else:
            for label, text in items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.leftIndent = Cm(0.6)
                p.paragraph_format.space_after = Pt(2)
                if label:
                    r1 = p.add_run(f"{label}：")
                    _set_fonts(r1, size=8.5, bold=True, color=NAVY)
                    r2 = p.add_run(text)
                    _set_fonts(r2, size=8.5, color=NAVY)
                else:
                    r = p.add_run(text)
                    _set_fonts(r, size=8.5, color=NAVY)

    # 升学目标
    _add_sub_heading(doc, "升学目标", "分阶段 · 可量化")
    if not target:
        _add_para(doc, "尚未填写升学目标。", size=8.5, color=MUTED)
    else:
        stages = _split_overview(target)
        # 识别短期/中期/高考三段
        if len(stages) <= 2 and "短期" not in target and "中期" not in target:
            _add_para(doc, target, size=8.5, color=NAVY)
        else:
            tbl = doc.add_table(rows=1, cols=3)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(tbl)
            # header
            labels = [s[0] or f"阶段{i+1}" for i, s in enumerate(stages[:3])]
            # 若不足3列，补齐
            while len(labels) < 3:
                labels.append("—")
            for ci, h in enumerate(labels[:3]):
                cell = tbl.cell(0, ci)
                _set_cell_shading(cell, NAVY_HEX)
                _set_cell_margins(cell, top=60, bottom=60)
                p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h); _set_fonts(run, size=8, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            row = tbl.add_row()
            for ci, (_, txt) in enumerate(stages[:3]):
                cell = row.cells[ci]
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                run = p.add_run(txt); _set_fonts(run, size=7.5, color=NAVY)
                p.paragraph_format.space_after = Pt(2)
            # 若超过3段，剩余用列表
            if len(stages) > 3:
                for label, txt in stages[3:]:
                    p = doc.add_paragraph(style="List Bullet")
                    r1 = p.add_run(f"{label}："); _set_fonts(r1, size=8, bold=True, color=NAVY)
                    r2 = p.add_run(txt); _set_fonts(r2, size=8, color=NAVY)

    # 当前状态
    _add_sub_heading(doc, "当前状态", f"负责人 · 班主任 #{case.owner_teacher_id}" if getattr(case, "owner_teacher_id", None) else "")
    _add_para(doc, summary or "尚未填写状态说明。", size=8.5, color=NAVY if summary else MUTED, space_before=2)


def _add_two_col_row(doc: Document, label: str, content: str, label_width=Cm(2.8)) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(tbl, color=BORDER_COLOR, size="4")
    c0 = tbl.cell(0, 0); c0.width = label_width
    _set_cell_shading(c0, HEADER_FILL); _set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
    c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label); _set_fonts(run, size=7.5, bold=True, color=MUTED)
    c1 = tbl.cell(0, 1)
    _set_cell_margins(c1, top=80, bottom=80)
    c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = c1.paragraphs[0]
    if not content or not content.strip():
        run = p.add_run("—"); _set_fonts(run, size=8, color=MUTED)
    else:
        # 按换行分段
        lines = content.strip().split("\n")
        for i, line in enumerate(lines):
            if i == 0:
                run = p.add_run(line.strip()); _set_fonts(run, size=8.5, color=NAVY)
            else:
                p2 = c1.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                run = p2.add_run(line.strip()); _set_fonts(run, size=8.5, color=NAVY)
    return tbl


def _add_subject_section(doc: Document, subject: str, plan: SubjectPlan | None, tasks: list[CaseTask], checkins: list[TaskCheckin], reviews: list[CaseReview], teacher_name: str = "") -> None:
    # 学科标题：大字 + 班底色
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "EFF6FF")
    pPr.append(shd)
    # 左侧色条模拟
    r_bar = p.add_run("  ▎  "); _set_fonts(r_bar, size=10, bold=True, color=ACCENT)
    r_title = p.add_run(subject)
    _set_fonts(r_title, size=12, bold=True, color=NAVY)
    if subject == "德育":
        r_sub = p.add_run("  ·  行为与规范")
    else:
        r_sub = p.add_run("  ·  学业提升方案")
    _set_fonts(r_sub, size=7.5, color=MUTED)
    if teacher_name:
        r_teacher = p.add_run(f"    责任教师：{teacher_name}")
        _set_fonts(r_teacher, size=7, color=MUTED)
    if tasks:
        r_cnt = p.add_run(f"    ·    {len(tasks)} 项任务")
        _set_fonts(r_cnt, size=7, color=MUTED)

    # 无方案提示
    if not plan:
        _add_para(doc, "该学科尚未建立方案。", size=8.5, color=MUTED, space_before=2)
        return

    # 诊断与归因
    _add_sub_heading(doc, "诊断与归因", "问题定位 · 原因剖析")
    _add_two_col_row(doc, "问题定位" if subject != "德育" else "行为表现", getattr(plan, "problem_location", "") or "")
    _add_two_col_row(doc, "原因剖析" if subject != "德育" else "原因/背景", getattr(plan, "cause_analysis", "") or "")

    # 目标与强化
    _add_sub_heading(doc, "目标与路径", "奋斗目标 · 高考/规范要求 · 具体强化")
    _add_two_col_row(doc, "奋斗目标" if subject != "德育" else "养成目标", getattr(plan, "struggle_goal", "") or "")
    _add_two_col_row(doc, "高考要求" if subject != "德育" else "规范要求", getattr(plan, "gaokao_requirement", "") or "")
    _add_two_col_row(doc, "具体强化" if subject != "德育" else "日常优化", getattr(plan, "reinforcement", "") or "")

    # 任务清单
    _add_sub_heading(doc, "任务清单", f"{len(tasks)} 项 · 日/周/月")
    if not tasks:
        _add_para(doc, "尚未安排任务。", size=8.5, color=MUTED)
    else:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(tbl)
        headers = ["#", "任务", "周期", "起止", "状态"]
        widths = [Cm(0.9), Cm(7.2), Cm(1.4), Cm(3.2), Cm(1.6)]
        for ci, h in enumerate(headers):
            cell = tbl.cell(0, ci)
            _set_cell_shading(cell, HEADER_FILL); _set_cell_margins(cell, top=50, bottom=50)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h); _set_fonts(run, size=7, bold=True, color=MUTED)
            cell.width = widths[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for idx, t in enumerate(tasks, 1):
            row = tbl.add_row()
            # 序号
            c = row.cells[0]; _set_cell_margins(c, top=50, bottom=50); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run(str(idx)); _set_fonts(run, size=7.5, color=MUTED)
            # 任务
            c = row.cells[1]; _set_cell_margins(c)
            p = c.paragraphs[0]; run = p.add_run(t.title or "—"); _set_fonts(run, size=8, bold=True, color=NAVY)
            if t.description:
                p2 = c.add_paragraph(); p2.paragraph_format.space_before = Pt(1); run = p2.add_run(t.description[:220]); _set_fonts(run, size=7, color=MUTED)
            # 周期
            c = row.cells[2]; _set_cell_margins(c); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run(_cadence_label(t.cadence)); _set_fonts(run, size=7.5, color=NAVY)
            # 起止
            c = row.cells[3]; _set_cell_margins(c); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run(f"{_format_date(t.starts_on)}  至  {_format_date(t.due_on)}"); _set_fonts(run, size=7, color=NAVY)
            # 状态
            c = row.cells[4]; _set_cell_margins(c); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            st_map = {"pending":"待开始","in_progress":"进行中","completed":"已完成","cancelled":"已取消"}
            run = p.add_run(st_map.get(t.status, t.status or "—")); _set_fonts(run, size=7, color=MUTED)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 表格整体宽度
        tbl.autofit = False

    # 执行记录
    rel = [ch for ch in (checkins or []) if any(ts.id == ch.task_id for ts in tasks)]
    if rel:
        _add_sub_heading(doc, "执行记录", f"{len(rel)} 条 · 班主任核实")
        for ch in sorted(rel, key=lambda x: x.checked_in_at, reverse=True)[:6]:
            # 找到任务标题
            t_title = next((t.title for t in tasks if t.id == ch.task_id), f"任务#{ch.task_id}")
            p = doc.add_paragraph()
            p.paragraph_format.leftIndent = Cm(0.4)
            p.paragraph_format.space_after = Pt(3)
            # 左侧完成度徽章模拟
            r_rate = p.add_run(f"[{ch.completion_rate}%]  ")
            color = RGBColor(0x16, 0xA3, 0x4A) if ch.completion_rate >= 80 else (RGBColor(0xD9, 0x77, 0x06) if ch.completion_rate >= 50 else MUTED)
            _set_fonts(r_rate, size=8, bold=True, color=color)
            r_t = p.add_run(f"{t_title}  ·  {_format_date(ch.checked_in_at)}")
            _set_fonts(r_t, size=7.5, bold=True, color=NAVY)
            if ch.self_check:
                p2 = doc.add_paragraph()
                p2.paragraph_format.leftIndent = Cm(0.8)
                p2.paragraph_format.space_after = Pt(4)
                run = p2.add_run(ch.self_check[:400]); _set_fonts(run, size=7.5, color=MUTED)

    # 督查
    cand = [r for r in (reviews or []) if not r.subject or r.subject == subject or (subject and subject in (r.subject or ""))]
    if not cand:
        cand = [r for r in (reviews or []) if r.review_level in {"head_teacher","subject"}] if subject != "德育" else []
    # 校级另取
    school_cand = [r for r in (reviews or []) if r.review_level in {"school","principal"} and (not r.subject or r.subject == subject)]
    show_reviews = cand[:2] + school_cand[:1]
    if show_reviews:
        _add_sub_heading(doc, "督查与复盘", "教师/校级")
        for r in sorted(show_reviews, key=lambda x: x.reviewed_at, reverse=True):
            level_map = {"head_teacher":"班主任","subject":"学科","school":"校级","principal":"校长","deyu":"德育"}
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r_lv = p.add_run(f"[{level_map.get(r.review_level, r.review_level)}]")
            _set_fonts(r_lv, size=7, bold=True, color=ACCENT)
            if r.subject:
                rs = p.add_run(f"  {r.subject}"); _set_fonts(rs, size=7, color=MUTED)
            rd = p.add_run(f"  {_format_date(r.reviewed_at)}")
            _set_fonts(rd, size=7, color=MUTED)
            if r.correction_due_on:
                rcd = p.add_run(f"  截止 {_format_date(r.correction_due_on)}")
                _set_fonts(rcd, size=7, color=RGBColor(0xBE, 0x12, 0x3C))
            if r.problem:
                p2 = doc.add_paragraph(); p2.paragraph_format.leftIndent = Cm(0.4); run = p2.add_run(f"问题：{r.problem[:300]}"); _set_fonts(run, size=7.5, color=NAVY)
            if r.corrective_action:
                p3 = doc.add_paragraph(); p3.paragraph_format.leftIndent = Cm(0.4); run = p3.add_run(f"整改：{r.corrective_action[:300]}"); _set_fonts(run, size=7.5, color=NAVY)
            if r.recheck_result:
                p4 = doc.add_paragraph(); p4.paragraph_format.leftIndent = Cm(0.4); run = p4.add_run(f"复查：{r.recheck_result[:200]}"); _set_fonts(run, size=7.5, color=MUTED)


def _template_cell_text(cell, value: object) -> None:
    """替换模板单元格文字，同时保留原段落及首个文字的版式。"""
    text = str(value or "—")
    paragraph = cell.paragraphs[0]
    saved_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            saved_rpr = deepcopy(run._r.rPr)
            break
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    new_run = paragraph.add_run(text)
    if saved_rpr is not None:
        new_run._r.insert(0, saved_rpr)
    paragraph.paragraph_format.space_after = Pt(0)


def _template_paragraph_text(paragraph, value: object) -> None:
    """替换模板段落文字，同时保留标题段落的字体和对齐样式。"""
    text = str(value or "—")
    saved_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            saved_rpr = deepcopy(run._r.rPr)
            break
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    new_run = paragraph.add_run(text)
    if saved_rpr is not None:
        new_run._r.insert(0, saved_rpr)


def _profile_value(profile: Any | None, key: str, default: str = "—") -> str:
    value = _pv(profile, key)
    if value is None:
        return default
    value = str(value).strip()
    return value or default


def _page_break_element() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    page_break = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break)
    paragraph.append(p_pr)
    return paragraph


def _paragraph_element(doc: Document, marker: str) -> OxmlElement | None:
    for paragraph in doc.paragraphs:
        if marker in (paragraph.text or ""):
            return paragraph._p
    return None


def _clone_subject_template_pages(doc: Document, count: int) -> Document:
    """按模板原样复制“每科一表”和任务表；克隆后重开以获得新的表格包装对象。"""
    if count <= 1:
        return doc
    subject_heading = _paragraph_element(doc, "学科方案（每科一表）")
    subject_title = _paragraph_element(doc, "科目 一生一案")
    task_heading = _paragraph_element(doc, "3.1 本学科任务")
    following_text = _paragraph_element(doc, "如有多学科方案")
    review_heading = _paragraph_element(doc, "督查、整改与阶段复盘")
    if any(element is None for element in (subject_heading, subject_title, task_heading, following_text, review_heading)) or len(doc.tables) < 4:
        raise ValueError("导出模板缺少学科方案复制所需的固定区域")

    source_elements = [subject_heading, subject_title, doc.tables[2]._tbl, task_heading, doc.tables[3]._tbl, following_text]
    # 依次在督查章节前插入副本，确保所有学科都保持模板的固定顺序与页面结构。
    for _ in range(count - 1):
        review_heading.addprevious(_page_break_element())
        for element in source_elements:
            clone = deepcopy(element)
            if element is subject_heading:
                # 模板首个学科标题自带分页；副本前已插入分页，须移除重复属性，避免空白页。
                p_pr = clone.find(qn("w:pPr"))
                if p_pr is not None:
                    page_break = p_pr.find(qn("w:pageBreakBefore"))
                    if page_break is not None:
                        p_pr.remove(page_break)
            review_heading.addprevious(clone)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return Document(bio)


def _fill_template_profile(table, profile: Any | None, student_name: str, class_name: str) -> None:
    """模板第 1 表：学生基本信息。模板结构固定，按行列填值而不修改单元格结构。"""
    if len(table.rows) < 9:
        raise ValueError("导出模板的学生基本信息表结构不完整")
    data = {
        (0, 1): student_name or _profile_value(profile, "student_name"),
        (0, 3): _profile_value(profile, "gender"),
        (1, 1): _profile_value(profile, "ethnicity"),
        (1, 3): _profile_value(profile, "grade", "高三"),
        (2, 1): class_name or "—",
        (2, 3): _profile_value(profile, "source_school"),
        (3, 1): _profile_value(profile, "parent_name"),
        (3, 3): _profile_value(profile, "parent_phone"),
        (7, 1): _profile_value(profile, "parent_evaluation"),
        (8, 1): _profile_value(profile, "primary_needs"),
    }
    for (row, col), value in data.items():
        _template_cell_text(table.cell(row, col), value)
    scores = "  ".join([
        f"语文 {_profile_value(profile, 'entrance_chinese', '—')}",
        f"数学 {_profile_value(profile, 'entrance_math', '—')}",
        f"英语 {_profile_value(profile, 'entrance_english', '—')}",
        f"物理 {_profile_value(profile, 'entrance_physics', '—')}",
        f"化学 {_profile_value(profile, 'entrance_chemistry', '—')}",
        f"生物 {_profile_value(profile, 'entrance_biology', '—')}",
        f"政治 {_profile_value(profile, 'entrance_politics', '—')}",
        f"历史 {_profile_value(profile, 'entrance_history', '—')}",
        f"地理 {_profile_value(profile, 'entrance_geography', '—')}",
        f"总分 {_profile_value(profile, 'entrance_total_score', '—')}",
    ])
    _template_cell_text(table.cell(6, 1), scores)
    if _pv(profile, "health_visible") is False:
        health = "健康信息已设为仅校长可见，本次导出不展示具体内容。"
        for row in (4, 5):
            _template_cell_text(table.cell(row, 1), health)
    else:
        _template_cell_text(table.cell(4, 1), _profile_value(profile, "other_health_notes"))
        _template_cell_text(table.cell(4, 3), _profile_value(profile, "underlying_conditions"))
        _template_cell_text(table.cell(5, 1), _profile_value(profile, "allergy_history"))


def _fill_template_subject(plan_table, task_table, plan: SubjectPlan, tasks: list[CaseTask], checkins: list[TaskCheckin], teacher_name: str) -> None:
    if len(plan_table.rows) < 6 or len(task_table.rows) < 2:
        raise ValueError("导出模板的学科方案区域不完整")
    subject = plan.subject or "—"
    fields = {
        (0, 1): subject,
        (0, 3): teacher_name or "—",
        (1, 1): plan.problem_location or "—",
        (2, 1): plan.cause_analysis or "—",
        (3, 1): plan.struggle_goal or "—",
        (4, 1): plan.gaokao_requirement or "—",
        (5, 1): plan.reinforcement or "—",
    }
    for (row, col), value in fields.items():
        _template_cell_text(plan_table.cell(row, col), value)
    latest_checkins = {}
    for checkin in checkins or []:
        if checkin.task_id not in latest_checkins or checkin.checked_in_at > latest_checkins[checkin.task_id].checked_in_at:
            latest_checkins[checkin.task_id] = checkin
    for row_index, task in enumerate(tasks[: len(task_table.rows) - 1], start=1):
        checkin = latest_checkins.get(task.id)
        _template_cell_text(task_table.cell(row_index, 0), task.title or "—")
        _template_cell_text(task_table.cell(row_index, 1), _format_date(task.due_on))
        _template_cell_text(task_table.cell(row_index, 2), f"{checkin.completion_rate}%" if checkin else "未打卡")
        _template_cell_text(task_table.cell(row_index, 3), (checkin.self_check if checkin and checkin.self_check else "待教师确认"))
        _template_cell_text(task_table.cell(row_index, 4), _format_date(checkin.checked_in_at) if checkin else "—")


def _fill_template_subject_titles(doc: Document, plans: list[SubjectPlan]) -> None:
    """将模板固定标题“科目一生一案”替换成各学科名称，避免导出后仍显示占位词。"""
    title_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if "科目" in (paragraph.text or "") and "一生一案" in (paragraph.text or "")
    ]
    for paragraph, plan in zip(title_paragraphs, plans):
        _template_paragraph_text(paragraph, f"{plan.subject or '—'}一生一案")


def _build_template_export_bytes(
    case: StudentCase,
    student_name: str,
    class_name: str,
    subject_plans: list[SubjectPlan],
    tasks: list[CaseTask],
    checkins: list[TaskCheckin],
    reviews: list[CaseReview],
    *,
    profile: Any | None,
    goals: list | None,
    teacher_names: dict[int, str] | None,
) -> bytes:
    """基于用户确认的 DOCX 底稿填充实际一生一案数据。"""
    doc = Document(CASE_EXPORT_TEMPLATE)
    _prepend_cover_image(doc)
    plan_map = {plan.subject: plan for plan in (subject_plans or [])}
    ordered_subjects = [subject for subject in SUBJECT_ORDER if subject in plan_map]
    ordered_subjects.extend(subject for subject in plan_map if subject not in ordered_subjects)
    ordered_plans = [plan_map[subject] for subject in ordered_subjects]
    # 无学科方案时仍保留模板自带的一张可填写学科页；有多科时按科复制。
    template_subject_count = max(1, len(ordered_plans))
    doc = _clone_subject_template_pages(doc, template_subject_count)
    if len(doc.tables) < 7:
        raise ValueError("导出模板表格数量不足")

    _fill_template_profile(doc.tables[0], profile, student_name, class_name)
    overview = doc.tables[1]
    _template_cell_text(overview.cell(0, 1), (case.overall_problem or "—") + (f"\n当前学习状态：{case.current_summary}" if case.current_summary else ""))
    _template_cell_text(overview.cell(1, 1), case.admission_target or "—")

    subject_table_start = 2
    _fill_template_subject_titles(doc, ordered_plans)
    for index, plan in enumerate(ordered_plans):
        plan_table = doc.tables[subject_table_start + index * 2]
        task_table = doc.tables[subject_table_start + index * 2 + 1]
        plan_tasks = sorted([task for task in (tasks or []) if task.subject == plan.subject], key=lambda task: task.due_on)
        teacher_name = (teacher_names or {}).get(getattr(plan, "teacher_id", None), "")
        _fill_template_subject(plan_table, task_table, plan, plan_tasks, checkins or [], teacher_name)

    after_subject_tables = subject_table_start + template_subject_count * 2
    review_table, adjustment_table, workflow_table = doc.tables[after_subject_tables: after_subject_tables + 3]
    for index, review in enumerate(sorted(reviews or [], key=lambda item: item.reviewed_at, reverse=True)[: len(review_table.rows) - 1], start=1):
        values = [
            {"head_teacher": "班主任", "subject": "学科", "deyu": "德育", "school": "校级", "principal": "校长"}.get(review.review_level, review.review_level),
            review.subject or "综合",
            review.problem or "—",
            review.corrective_action or "—",
            _format_date(review.correction_due_on),
            review.recheck_result or "待复查",
            "—",
            _format_date(review.reviewed_at),
        ]
        for col, value in enumerate(values):
            _template_cell_text(review_table.cell(index, col), value)
    if goals:
        goal_summary = "；".join(f"{goal.title}：{goal.baseline_value or '—'}→{goal.target_value or '—'}" for goal in goals[:3])
        _template_cell_text(adjustment_table.cell(0, 3), goal_summary or "—")
    _template_cell_text(adjustment_table.cell(2, 1), f"V{case.version}")
    _template_cell_text(adjustment_table.cell(2, 3), _status_label(case.status))
    _template_cell_text(workflow_table.cell(1, 1), _status_label(case.status))
    _template_cell_text(workflow_table.cell(1, 4), _format_date(case.updated_at))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_case_export_bytes(
    case: StudentCase,
    student_name: str,
    class_name: str,
    cycle_name: str,
    subject_plans: list[SubjectPlan],
    tasks: list[CaseTask],
    checkins: list[TaskCheckin],
    reviews: list[CaseReview],
    cycle: CaseCycle | None = None,
    *,
    profile: Any | None = None,
    goals: list | None = None,
    guardians: list | None = None,
    teacher_names: dict[int, str] | None = None,
) -> bytes:
    """生成正式版式 DOCX（A4、分节、可打印）。

    兼容旧调用：profile/goals 可为空，内部降级。
    """
    if CASE_EXPORT_TEMPLATE.exists():
        return _build_template_export_bytes(
            case=case,
            student_name=student_name,
            class_name=class_name,
            subject_plans=subject_plans,
            tasks=tasks,
            checkins=checkins,
            reviews=reviews,
            profile=profile,
            goals=goals,
            teacher_names=teacher_names,
        )

    doc = Document()
    _prepend_cover_image(doc)
    _ensure_a4_margins(doc)
    _add_header_footer(doc, cycle_name=cycle_name)

    # 样式：正文默认
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.get_or_add_rPr().find(qn("w:rFonts")).set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(9)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    _add_cover(doc, student_name, class_name, cycle_name, case, profile=profile)

    # 档案信息
    _add_profile_section(doc, profile, guardians=guardians)

    # 总览
    _add_overview_section(doc, case)

    # 学科方案
    _add_section_heading(doc, "03", "学科方案", f"{len(subject_plans or [])} 科 · 含任务与督查")
    if not subject_plans:
        _add_para(doc, "尚未建立学科方案。", size=8.5, color=MUTED)
    else:
        plan_map = {p.subject: p for p in subject_plans}
        ordered: list[str] = []
        for s in SUBJECT_ORDER:
            if s in plan_map:
                ordered.append(s)
        for s in plan_map:
            if s not in ordered:
                ordered.append(s)
        for idx, subject in enumerate(ordered):
            plan = plan_map.get(subject)
            subj_tasks = [t for t in (tasks or []) if t.subject == subject]
            subj_tasks = sorted(subj_tasks, key=lambda x: x.due_on)
            # 学科间分页：首科不分页，后续每科前分页
            if idx > 0:
                p = doc.add_paragraph()
                p.paragraph_format.pageBreakBefore = True
                # 空段用于分页标记
                run = p.add_run("")
                _set_fonts(run, size=1, color=RGBColor(0xFF, 0xFF, 0xFF))
            teacher_name = ""
            if plan and teacher_names and getattr(plan, "teacher_id", None) in teacher_names:
                teacher_name = teacher_names[plan.teacher_id]
            _add_subject_section(doc, subject, plan, subj_tasks, checkins or [], reviews or [], teacher_name=teacher_name)

    # 附录
    sec = doc.add_paragraph()
    sec.paragraph_format.pageBreakBefore = True
    run = sec.add_run("")
    _set_fonts(run, size=1, color=RGBColor(0xFF, 0xFF, 0xFF))
    _add_section_heading(doc, "附录", "说明与版本", "Appendix")
    _add_para(doc, "填写与流转说明", size=9, bold=True, color=NAVY, space_before=4, space_after=2)
    bullets = [
        "一人一科一表：综合材料须按学科拆分为独立记录，杜绝合并标签；已存在学科专属文档以专属内容为准。",
        "证据锚定：问题定位尽量标注题号/得分/日期（如“文言翻译题3-2 扣4分”），便于与作业/考试证据关联。",
        "德育单列：作为独立学科记录课堂/作业/纪律/品行，用“养成目标/规范要求/日常优化”替代高考要求，保持同一底表。",
        "流转：草稿 → 待德育审查 → 执行中 → 待复盘 → 已调整 → 已归档；执行中正文锁定，调整须发起复盘并生成新版本。",
        "家长仅可见已确认版本；健康与体检信息可设为仅校长可见。",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.leftIndent = Cm(0.6)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(b); _set_fonts(run, size=7.5, color=MUTED)

    _add_para(doc, f"版本信息：V{case.version}  ·  {_status_label(case.status)}  ·  创建 {_format_date(case.created_at)}  ·  更新 {_format_date(case.updated_at)}", size=7, color=MUTED, space_before=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "白粉高级英才中学  ·  高三一生一案工作组", size=7, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=1)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """使用本机办公转换器把已排版 DOCX 转成 PDF；无转换器时明确失败。"""
    with tempfile.TemporaryDirectory(prefix="case-export-") as temp_dir:
        temp = Path(temp_dir)
        source = temp / "case.docx"
        source.write_bytes(docx_bytes)
        output = temp / "case.pdf"
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            lo_profile = temp / "lo-profile"
            result = subprocess.run(
                [soffice, "--headless", f"-env:UserInstallation=file:///{lo_profile.as_posix()}", "--convert-to", "pdf", "--outdir", str(temp), str(source)],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if result.returncode == 0 and output.exists():
                return output.read_bytes()
        # Windows 部署常见为已安装 Microsoft Word 但未安装 LibreOffice。
        if shutil.which("powershell.exe"):
            script = (
                "$inputPath=$args[0]; $outputPath=$args[1]; "
                "$word=New-Object -ComObject Word.Application; $word.Visible=$false; "
                "try { $doc=$word.Documents.Open($inputPath, $false, $true); "
                "$doc.SaveAs2($outputPath, 17); $doc.Close($false) } "
                "finally { $word.Quit() }"
            )
            result = subprocess.run(
                ["powershell.exe", "-NoProfile", "-NonInteractive", "-Command", script, str(source), str(output)],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            if result.returncode == 0 and output.exists():
                return output.read_bytes()
        raise RuntimeError("PDF 导出需要 LibreOffice 或 Microsoft Word 转换器，请在服务环境安装其一")


def build_case_export_pdf_bytes(*args, **kwargs) -> bytes:
    """生成学校格式 PDF；底层仍复用已验收的 DOCX 模板填充逻辑。"""
    return _convert_docx_to_pdf(build_case_export_bytes(*args, **kwargs))


def _register_pdf_font() -> str:
    """选择部署环境已有的中文字体；找不到时退回 Helvetica，避免接口启动失败。"""
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            name = "CaseExportCJK"
            if name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, str(font_path), subfontIndex=0))
            return name
    return "Helvetica"


def _pdf_text(value: Any, fallback: str = "—") -> str:
    text = "" if value is None else str(value).strip()
    return escape(text or fallback).replace("\n", "<br/>")


def _pdf_value(obj: Any, name: str, fallback: str = "—") -> str:
    return _pdf_text(getattr(obj, name, None), fallback)


def _pdf_paragraph(value: Any, style: ParagraphStyle, fallback: str = "—") -> Paragraph:
    return Paragraph(_pdf_text(value, fallback), style)


def _pdf_table(rows: list[list[Any]], widths: list[float], styles: dict[str, ParagraphStyle], *, header: bool = True) -> Table:
    converted: list[list[Any]] = []
    for row_index, row in enumerate(rows):
        converted.append([
            cell if isinstance(cell, (Paragraph, Image)) else _pdf_paragraph(cell, styles["table_header" if header and row_index == 0 else "table"])
            for cell in row
        ])
    table = Table(converted, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#E2E8F0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if header:
        commands.extend([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ])
    table.setStyle(TableStyle(commands))
    return table


def _native_pdf_page(canvas, doc, font_name: str) -> None:
    canvas.saveState()
    canvas.setFont(font_name, 8)
    canvas.setFillColor(colors.HexColor("#64748B"))
    canvas.drawString(2 * cm, 1.15 * cm, "白粉高级英才中学 · 高三一生一案工作组")
    canvas.drawRightString(A4[0] - 2 * cm, 1.15 * cm, f"第 {doc.page} 页")
    canvas.restoreState()


def _build_native_pdf_bytes(
    case: StudentCase,
    student_name: str,
    class_name: str,
    cycle_name: str,
    subject_plans: list[SubjectPlan],
    tasks: list[CaseTask],
    checkins: list[TaskCheckin],
    reviews: list[CaseReview],
    cycle: CaseCycle | None = None,
    *,
    profile: Any | None = None,
    goals: list | None = None,
    guardians: list | None = None,
    teacher_names: dict[int, str] | None = None,
) -> bytes:
    """按现有 DOCX 导出内容生成原生 PDF；所有内容在内存中完成。"""
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    styles_map = {
        "title": ParagraphStyle("case-title", parent=styles["Title"], fontName=font_name, fontSize=22, leading=30, textColor=colors.HexColor("#1B2F4A"), alignment=TA_CENTER, spaceAfter=18),
        "subtitle": ParagraphStyle("case-subtitle", parent=styles["Normal"], fontName=font_name, fontSize=11, leading=18, textColor=colors.HexColor("#64748B"), alignment=TA_CENTER),
        "section": ParagraphStyle("case-section", parent=styles["Heading1"], fontName=font_name, fontSize=15, leading=22, textColor=colors.HexColor("#1B2F4A"), spaceBefore=8, spaceAfter=8),
        "subject": ParagraphStyle("case-subject", parent=styles["Heading2"], fontName=font_name, fontSize=13, leading=19, textColor=colors.HexColor("#2F6FED"), spaceBefore=7, spaceAfter=6),
        "body": ParagraphStyle("case-body", parent=styles["BodyText"], fontName=font_name, fontSize=9, leading=14, textColor=colors.HexColor("#334155"), spaceAfter=4),
        "table": ParagraphStyle("case-table", parent=styles["BodyText"], fontName=font_name, fontSize=8.5, leading=12, textColor=colors.HexColor("#334155")),
        "table_header": ParagraphStyle("case-table-header", parent=styles["BodyText"], fontName=font_name, fontSize=8.5, leading=12, textColor=colors.white),
        "small": ParagraphStyle("case-small", parent=styles["BodyText"], fontName=font_name, fontSize=7.5, leading=11, textColor=colors.HexColor("#64748B")),
    }
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=1.8 * cm, bottomMargin=1.8 * cm, title=f"{student_name} 一生一案")
    story: list[Any] = [Spacer(1, 2.2 * cm)]
    if COVER_IMAGE.exists():
        story.extend([Image(str(COVER_IMAGE), width=15.8 * cm, height=8.8 * cm), Spacer(1, 1.2 * cm)])
    story.extend([
        Paragraph("个性化学习发展指导方案", styles_map["title"]),
        Paragraph(_pdf_text(student_name), styles_map["subtitle"]),
        Paragraph(f"{_pdf_text(class_name)} · {_pdf_text(cycle_name)} · V{getattr(case, 'version', 1)} · {_pdf_text(_status_label(getattr(case, 'status', '')))}", styles_map["subtitle"]),
        PageBreak(),
        Paragraph("01  学生基本信息", styles_map["section"]),
    ])
    profile_rows = [["项目", "内容", "项目", "内容"]]
    profile = profile or object()
    profile_rows.extend([
        ["姓名", _pdf_value(profile, "student_name", student_name), "性别", _pdf_value(profile, "gender")],
        ["年级", _pdf_value(profile, "grade"), "来源学校", _pdf_value(profile, "source_school")],
        ["家长", _pdf_value(profile, "parent_name"), "联系电话", _pdf_value(profile, "parent_phone")],
        ["主要需求", _pdf_value(profile, "primary_needs"), "家长评价", _pdf_value(profile, "parent_evaluation")],
    ])
    story.append(_pdf_table(profile_rows, [2.2 * cm, 6.3 * cm, 2.2 * cm, 6.3 * cm], styles_map))
    score_fields = [("语文", "entrance_chinese"), ("数学", "entrance_math"), ("英语", "entrance_english"), ("物理", "entrance_physics"), ("化学", "entrance_chemistry"), ("生物", "entrance_biology"), ("政治", "entrance_politics"), ("历史", "entrance_history"), ("地理", "entrance_geography")]
    story.extend([Spacer(1, 8), Paragraph("入学成绩", styles_map["subject"]), _pdf_table([["科目", "分数"]] + [[subject, _pdf_value(profile, field)] for subject, field in score_fields], [4 * cm, 4 * cm], styles_map)])
    story.extend([Spacer(1, 8), Paragraph("02  总案概览", styles_map["section"]), _pdf_table([
        ["总体问题", _pdf_value(case, "overall_problem")],
        ["升学目标", _pdf_value(case, "admission_target")],
        ["当前状态", _pdf_value(case, "current_summary")],
    ], [3.2 * cm, 14 * cm], styles_map, header=False)])
    if goals:
        story.extend([Spacer(1, 8), Paragraph("阶段目标", styles_map["subject"]), _pdf_table([["目标", "基线", "目标值"]] + [[getattr(g, "title", ""), getattr(g, "baseline_value", "") or "—", getattr(g, "target_value", "") or "—"] for g in goals], [7 * cm, 5 * cm, 5.2 * cm], styles_map)])
    story.extend([PageBreak(), Paragraph("03  学科方案", styles_map["section"])])
    plan_map = {p.subject: p for p in (subject_plans or [])}
    ordered = [s for s in SUBJECT_ORDER if s in plan_map] + [s for s in plan_map if s not in SUBJECT_ORDER]
    for index, subject in enumerate(ordered):
        plan = plan_map[subject]
        if index:
            story.append(PageBreak())
        teacher = (teacher_names or {}).get(getattr(plan, "teacher_id", None), "")
        story.extend([Paragraph(f"{subject}一生一案", styles_map["subject"]), _pdf_table([
            ["负责教师", teacher or "—"], ["问题定位", getattr(plan, "problem_location", "")], ["原因分析", getattr(plan, "cause_analysis", "")],
            ["攻坚目标", getattr(plan, "struggle_goal", "")], ["高考要求", getattr(plan, "gaokao_requirement", "")], ["强化安排", getattr(plan, "reinforcement", "")],
        ], [3.2 * cm, 14 * cm], styles_map, header=False)])
        subject_tasks = sorted([t for t in (tasks or []) if getattr(t, "subject", None) == subject], key=lambda t: getattr(t, "due_on", None) or "")
        if subject_tasks:
            task_rows = [["任务", "截止日期", "完成率", "打卡记录"]]
            checkin_map = {getattr(c, "task_id", None): c for c in (checkins or [])}
            for task in subject_tasks:
                checkin = checkin_map.get(getattr(task, "id", None))
                task_rows.append([getattr(task, "title", ""), _format_date(getattr(task, "due_on", None)), f"{getattr(checkin, 'completion_rate', '—')}%" if checkin else "—", getattr(checkin, "self_check", "") if checkin else "未打卡"])
            story.extend([Spacer(1, 6), Paragraph("任务与执行", styles_map["body"]), _pdf_table(task_rows, [5.2 * cm, 3.2 * cm, 2.6 * cm, 6.2 * cm], styles_map)])
    if not ordered:
        story.append(Paragraph("尚未建立学科方案。", styles_map["body"]))
    story.extend([PageBreak(), Paragraph("04  版本与说明", styles_map["section"]), Paragraph("家长仅可见已确认版本；健康与体检信息按当前权限控制，不在无权导出内容中展示。", styles_map["body"]), Paragraph(f"版本信息：V{getattr(case, 'version', 1)} · {_pdf_text(_status_label(getattr(case, 'status', '')))} · 创建 {_pdf_text(_format_date(getattr(case, 'created_at', None)))} · 更新 {_pdf_text(_format_date(getattr(case, 'updated_at', None)))}", styles_map["small"])])
    doc.build(story, onFirstPage=lambda c, d: _native_pdf_page(c, d, font_name), onLaterPages=lambda c, d: _native_pdf_page(c, d, font_name))
    return buffer.getvalue()
