from datetime import date, datetime, timezone
from io import BytesIO
from types import SimpleNamespace

from docx import Document

from app.services.case_export import build_case_export_bytes


def test_export_uses_configured_template_and_repeats_subject_pages():
    """正式导出应以用户确认的底稿为基础，且每科都有独立的模板页。"""
    now = datetime(2026, 9, 1, tzinfo=timezone.utc)
    case = SimpleNamespace(
        overall_problem="函数与阅读理解存在短板",
        admission_target="一本院校",
        current_summary="执行第一阶段计划",
        version=2,
        status="executing",
        updated_at=now,
        created_at=now,
    )
    profile = SimpleNamespace(
        student_name="测试学生", gender="女", ethnicity="汉族", grade="高三",
        source_school="测试中学", parent_name="测试家长", parent_phone="13800138000",
        parent_evaluation="学习态度认真", primary_needs="提升时间管理", health_visible=True,
        other_health_notes="无", underlying_conditions="无", allergy_history="无",
        entrance_chinese=105, entrance_math=98, entrance_english=110, entrance_physics=82,
        entrance_chemistry=76, entrance_biology=80, entrance_politics=85,
        entrance_history=None, entrance_geography=None, entrance_total_score=636,
    )
    plans = [
        SimpleNamespace(subject="数学", teacher_id=2, problem_location="函数综合题失分", cause_analysis="审题不完整", struggle_goal="函数得分提升", gaokao_requirement="掌握导数应用", reinforcement="每日一题"),
        SimpleNamespace(subject="英语", teacher_id=3, problem_location="阅读速度偏慢", cause_analysis="词汇量不足", struggle_goal="阅读正确率提升", gaokao_requirement="阅读理解", reinforcement="每日精读"),
    ]
    tasks = [
        SimpleNamespace(id=1, subject="数学", title="函数专题训练", due_on=date(2026, 9, 10)),
        SimpleNamespace(id=2, subject="英语", title="阅读精练", due_on=date(2026, 9, 12)),
    ]
    checkins = [SimpleNamespace(task_id=1, checked_in_at=now, completion_rate=80, self_check="已完成基础题")]
    doc = Document(BytesIO(build_case_export_bytes(
        case, "测试学生", "高三1班", "第一阶段", plans, tasks, checkins, [],
        profile=profile, goals=[], teacher_names={2: "张老师", 3: "李老师"},
    )))
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "01学生基本信息" in text
    assert "数学一生一案" in text
    assert "英语一生一案" in text
    assert "科目一生一案" not in text
    assert "测试学生" in table_text
    assert "数学" in table_text
    assert "英语" in table_text
    assert "函数专题训练" in table_text
    assert "阅读精练" in table_text
    assert len(doc.sections) >= 2
    assert doc.sections[0].left_margin == 0
    assert doc.inline_shapes[0].width >= 21 * 360000
    assert doc.inline_shapes[0].height >= 29.7 * 360000
