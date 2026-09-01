"""Build a black-and-white editable Word template for the current student-case workflow."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("deliverables") / "一生一案_导出模板_可编辑版.docx"
BLACK = "000000"
WHITE = "FFFFFF"


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = __import__("docx").shared.RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:hint"), "eastAsia")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, bottom=90, left=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "8")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), BLACK)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_fixed_widths(table, widths_cm):
    table.autofit = False
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_cm):
        # w:gridCol/@w:w uses twips, not EMU.  Writing EMU here made Word
        # interpret every table as hundreds of inches wide and shrink it.
        col.set(qn("w:w"), str(Cm(width).twips))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


def add_text(cell, text, bold=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    return p


def add_heading(doc, title, level=1, subtitle=""):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run_font(r, size=16 if level == 1 else 14, bold=True, color=BLACK)
    if subtitle:
        r = p.add_run(f"  {subtitle}")
        set_run_font(r, size=10, color=BLACK)
    return p


def add_body(doc, text, indent=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_kv_table(doc, pairs, widths=(3.0, 6.3, 3.0, 6.3)):
    rows = (len(pairs) + 1) // 2
    table = doc.add_table(rows=rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_widths(table, widths)
    set_table_borders(table)
    for idx, (label, value) in enumerate(pairs):
        row, col = divmod(idx, 2)
        label_cell, value_cell = table.cell(row, col * 2), table.cell(row, col * 2 + 1)
        # 用户指定所有表格均为白底黑框：字段名区仅用加粗区分，不再使用深色底。
        set_cell_shading(label_cell, WHITE)
        add_text(label_cell, label, bold=True, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text(value_cell, value)
    if len(pairs) % 2:
        table.cell(rows - 1, 2).merge(table.cell(rows - 1, 3))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_grid_table(doc, headers, rows, widths, blank_rows=0):
    table = doc.add_table(rows=1 + len(rows) + blank_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_widths(table, widths)
    set_table_borders(table)
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, WHITE)
        add_text(cell, text, bold=True, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    repeat_header(table.rows[0])
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, len(row_data) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            add_text(table.cell(r_idx, c_idx), text, align=align, size=11)
    for r_idx in range(1 + len(rows), 1 + len(rows) + blank_rows):
        table.rows[r_idx].height = Cm(0.8)
        table.rows[r_idx].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c_idx in range(len(headers)):
            add_text(table.cell(r_idx, c_idx), "", size=11)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_student_basic_info_table(doc):
    """A single editable table; deliberately avoids merged cells for Word compatibility."""
    table = doc.add_table(rows=7, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_widths(table, [1.9, 4.1, 1.9, 4.1, 1.9, 4.1])
    set_table_borders(table)

    data = [
        ("姓名", "【填写】", "性别", "【填写】", "照片位", "建议粘贴一寸或两寸证件照"),
        ("民族", "【填写】", "年级", "高三", "照片说明", "可替换为学生证件照"),
        ("班级", "【填写】", "生源学校", "【填写】", "家长姓名", "【填写】"),
        ("联系电话", "【填写】", "与学生关系", "【填写】", "主要需求", "【填写】"),
        ("家长评价", "【填写学生学习习惯、成长表现、家庭支持情况】", "入学/基线总分", "【填写】", "健康信息权限", "【按系统 health_visible 设置】"),
        ("分科成绩", "语文【】 数学【】 英语【】 物理【】 化学【】 生物【】 政治【】 历史【】 地理【】", "过敏史", "【如无，填写“无”】", "基础疾病", "【如无，填写“无”】"),
        ("其他健康提示", "【填写】", "导出提示", "健康信息按实际权限填写/导出", "填表备注", "无对应信息时填写“无”或“—”"),
    ]
    for r, row in enumerate(data):
        for c, text in enumerate(row):
            label = c % 2 == 0
            add_text(table.cell(r, c), text, bold=label, align=WD_ALIGN_PARAGRAPH.CENTER if label else WD_ALIGN_PARAGRAPH.LEFT, size=10.5)
    table.rows[0].height = Cm(3.4)
    for row in table.rows[1:]: row.height = Cm(1.05)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_integrated_plan_table(doc):
    """One continuous editable table without merged cells."""
    table = doc.add_table(rows=14, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_widths(table, [1.4, 2.3, 2.7, 4.3, 4.3, 3.3])
    set_table_borders(table)

    rows = [
        ["总案概述", "填写内容", "总案概述", "填写内容", "总案概述", "填写内容"],
        ["总体问题", "【从学业、习惯、情绪、时间管理等维度概括】", "升学目标", "【目标院校/专业方向/成绩或位次目标】", "当前状态", "【当前优势、短板和阶段表现】"],
        ["综合研判", "【班主任填写】", "诊断数量", "【填写】", "证据状态", "【已确认/待确认】"],
        ["序号", "学科/维度", "问题类别", "诊断内容", "证据来源/定位", "教师确认"],
        ["1", "【填写】", "【知识/能力/习惯/德育等】", "【具体问题】", "【考试日期、题号、作业/反馈记录】", "【姓名/日期】"],
        ["", "", "", "", "", ""],
        ["目标与任务", "填写内容", "目标与任务", "填写内容", "目标与任务", "填写内容"],
        ["目标类型", "学科", "目标名称", "基线值/目标值", "截止日期", "状态"],
        ["【学业/习惯/阶段】", "【填写或综合】", "【明确、可衡量】", "【基线】→【目标】", "【YYYY-MM-DD】", "【未开始/进行中/已完成】"],
        ["", "", "", "", "", ""],
        ["任务名称", "学科", "任务说明", "频次", "起止日期", "责任人/状态"],
        ["【填写】", "【填写】", "【具体到材料、题型、时长或数量】", "【日/周/月】", "【开始】至【截止】", "【教师/学生】\n【待执行】"],
        ["", "", "", "", "", ""],
    ]
    header_rows = {0, 3, 6, 7, 10}
    for r, row_data in enumerate(rows):
        for c, text in enumerate(row_data):
            add_text(table.cell(r, c), text, bold=r in header_rows or (r in {1, 2, 14, 15} and c % 2 == 0), align=WD_ALIGN_PARAGRAPH.CENTER if r in header_rows or c == 0 else WD_ALIGN_PARAGRAPH.LEFT, size=10 if r in header_rows else 10.5)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for idx in (5, 9, 12, 13): table.rows[idx].height = Cm(0.35)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_page_number(footer):
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第 ")
    set_run_font(r, size=9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)
    r = p.add_run(" 页")
    set_run_font(r, size=9)


def main():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(2.0), Cm(2.0)
    section.left_margin, section.right_margin = Cm(2.0), Cm(2.0)
    section.header_distance, section.footer_distance = Cm(1.0), Cm(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for style_name, size, before, after in (("Heading 1", 16, 12, 5), ("Heading 2", 14, 8, 4)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.add_run("一生一案学业发展管理系统 · 高三试点"), size=9)
    set_page_number(section.footer)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("【学校名称】"), size=18, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("一 生 一 案"), size=26, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("个性化学习发展指导方案（导出模板）"), size=16)
    doc.add_paragraph()
    add_kv_table(doc, [
        ("学生姓名", "【填写】"), ("班级", "【填写】"),
        ("学年/周期", "【如：2026-2027学年 第一阶段】"), ("方案版本", "V【填写】"),
        ("方案状态", "【草稿/待德育审查/需整改/执行中/待复盘/已调整/已归档】"), ("建档日期", "【YYYY年MM月DD日】"),
        ("班主任", "【填写】"), ("最后更新", "【YYYY年MM月DD日】"),
    ])
    add_body(doc, "使用说明：本模板对应当前系统的“诊断证据—总案—学科方案—目标任务—执行记录—督查复盘—版本留痕”闭环。填写内容以教师确认的信息和系统证据为准。", bold_prefix="使用说明：")
    page_break(doc)

    # 三类资料归并为一张基本信息表，方便与证件照一起归档、单页打印。
    add_heading(doc, "01  学生基本信息", subtitle="Student Profile")
    add_student_basic_info_table(doc)
    page_break(doc)

    # 总案、诊断、目标和任务保留在同一张连续工作表中，便于纵向查看完整依据与行动。
    add_heading(doc, "02  总案、诊断证据与学业目标阶段任务", subtitle="Overall Plan, Evidence, Goals & Tasks")
    add_integrated_plan_table(doc)
    page_break(doc)

    add_heading(doc, "04  学科方案（每科一表）", subtitle="Subject Plan")
    add_body(doc, "填写规则：每个已建立方案的学科单独复制本节；以系统学科方案字段为准。德育单列，重点填写养成目标、规范要求与日常优化。", bold_prefix="填写规则：")
    add_kv_table(doc, [
        ("学科", "【语文/数学/英语/物理/化学/生物/政治/历史/地理/德育】"), ("任课教师", "【填写】"),
        ("问题定位", "【具体到知识点、题型、能力或行为表现；可关联诊断证据】"), ("原因分析", "【填写】"),
        ("奋斗目标", "【阶段性、可检验目标】"), ("高考要求/养成要求", "【学科填写高考要求；德育填写养成与规范要求】"),
        ("巩固强化措施", "【填写具体方法、资源与频次】"), ("方案状态", "【草稿/已确认等】"),
    ])
    add_heading(doc, "4.1  本学科任务与执行记录", level=2)
    add_grid_table(doc, ["任务", "计划完成时间", "完成率", "学生自查", "班主任/教师确认", "记录日期"], [
        ["【任务名称】", "【YYYY-MM-DD】", "【0-100%】", "【填写】", "【填写】", "【YYYY-MM-DD】"],
    ], [3.4, 3.0, 2.0, 4.0, 4.0, 2.6], blank_rows=2)
    add_body(doc, "如有多学科方案，请复制本节后继续填写。")

    # 督查表较宽，整节另起页，避免标题留在上一页而表格拆到下一页。
    page_break(doc)
    add_heading(doc, "05  督查、整改与阶段复盘", subtitle="Review & Adjustment")
    add_heading(doc, "5.1  督查记录", level=2)
    add_grid_table(doc, ["督查层级", "学科/维度", "发现问题", "整改措施", "整改截止日期", "复查结果", "督查人", "日期"], [
        ["【班主任/德育/校级/校长】", "【填写】", "【具体描述】", "【明确行动】", "【YYYY-MM-DD】", "【填写】", "【填写】", "【YYYY-MM-DD】"],
    ], [2.4, 2.0, 3.4, 3.4, 2.5, 3.2, 2.0, 2.3], blank_rows=3)
    add_heading(doc, "5.2  阶段复盘与版本调整", level=2)
    add_kv_table(doc, [
        ("复盘周期", "【填写：如第X周/第一次月考后】"), ("复盘结论", "【目标达成、问题变化、证据依据】"),
        ("调整原因", "【填写】"), ("调整内容", "【总案/学科方案/目标/任务的具体变更】"),
        ("新版本号", "V【填写】"), ("调整后状态", "【待德育审查/执行中/已调整等】"),
    ])

    # 流转表保持完整页，避免最后一行跨页造成阅读和填写断裂。
    page_break(doc)
    add_heading(doc, "06  审查、流转与版本留痕", subtitle="Workflow & Version")
    add_grid_table(doc, ["环节", "状态/结论", "责任角色", "意见或原因", "日期"], [
        ["班主任提交", "【待德育审查】", "班主任", "【填写】", "【YYYY-MM-DD】"],
        ["德育审查", "【通过/退回修改】", "德育主任", "【填写问题、整改要求；退回须写截止日期】", "【YYYY-MM-DD】"],
        ["进入执行", "【执行中】", "班主任/学科教师", "【填写】", "【YYYY-MM-DD】"],
        ["阶段复盘", "【待复盘/已调整】", "班主任/相关人员", "【填写】", "【YYYY-MM-DD】"],
        ["归档", "【已归档】", "管理人员", "【填写】", "【YYYY-MM-DD】"],
    ], [2.5, 3.0, 2.8, 7.0, 3.0])
    add_body(doc, "流转提醒：正式方案经德育审查后进入执行；执行中的正文不得无痕覆盖，需经阶段复盘形成新版本。家长仅查看系统已确认并具备可见权限的版本。")
    add_heading(doc, "07  填写与导出核对", subtitle="Checklist")
    for text in [
        "□ 学生档案、家长信息和基线成绩已核对；健康信息已按权限处理。",
        "□ 总体问题、升学目标和诊断证据已由教师确认，证据可追溯到考试、作业或反馈记录。",
        "□ 每个学科方案均包含问题定位、原因分析、目标、要求与强化措施。",
        "□ 目标、任务、执行记录、督查及复盘信息已与当前系统数据一致。",
        "□ 导出版本号、状态、日期和可见范围已核对；未将内部督查内容误发给无权限对象。",
    ]:
        add_body(doc, text)

    doc.core_properties.title = "一生一案导出模板（可编辑版）"
    doc.core_properties.subject = "高三一生一案学业发展管理系统"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
